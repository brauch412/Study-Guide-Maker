#!/usr/bin/env python
# -*- coding: utf-8 -*-
from flask import Flask
from flask import request
from flask import render_template
from flask import send_file
from flask import make_response
from docx import Document
from flask import redirect
from cStringIO import StringIO
from mstranslator import Translator
from wikiapi import WikiApi
from PyDictionary import PyDictionary
from wordnik import *
from textwrap import wrap
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from mstranslator import Translator
import re
import string
import os
import unicodedata
import urllib2
import shutil
import time
import os

from flask_cors import CORS, cross_origin


translator = Translator('sgm', 'byjuDVe1vZj8uyLtHiBZCBA4HDdoCQ8JZl5mfFvsZTM=')
wiki = WikiApi()
wiki = WikiApi({ 'locale' : 'en'})
app = Flask(__name__)
dictionary=PyDictionary()
apiUrl = 'http://api.wordnik.com/v4'
apiKey = 'c9d407740a1cd0f9af20d0886020975b631cc469845c39152'
client = swagger.ApiClient(apiKey, apiUrl)
wordApi = WordApi.WordApi(client)


"""
This flask app will be hosted using https://www.pythonanywhere.com
"""

@app.route('/tmp/<path:filename>', methods=['GET'])
@cross_origin()
def base_static(filename):
    response = make_response(send_file('static/temp/' + filename))
    response.headers['Content-Disposition'] = "attachment; filename=StudyGuideMaker.pdf"
    response.mimetype = 'application/pdf'
    return response
    #return send_file('static/temp/' + filename)

@app.route('/')
def main_page():
    return send_file('static/index.html')

@app.route('/originalsite')
def my_form():
    return render_template('index.html')


@app.route('/temp', methods = ['POST'])
def temp():

    term = request.form['term']
    #term = re.sub(u"\u2013", "-", term)
    #term = unicodedata.normalize('NFD', term).encode('ascii', 'ignore')
    term = term.encode('utf8')
    stringer = str(term)
    stringterm = os.linesep.join([s for s in stringer.splitlines() if s])

    if ('\n' in stringterm) and (',' in stringterm):
        clean = stringterm.replace(',','')
        cleansed = [s.strip() for s in clean.splitlines()]
    if ('\n' in stringterm) and (',' not in stringterm):
        cleansed = stringterm.splitlines()
    if ('\n' not in stringterm) and (','  in stringterm):
        cleansed = stringterm.split(',')
    if ('\n\n' in stringterm) and (',' not in stringterm):
        cleansed = stringterm.rstrip()

    #handle one term case
    try:
        arb=len(cleansed)
    except UnboundLocalError:
        cleansed=[stringterm]

    myarr = []
    output = StringIO()

    if cleansed[0] == "E2F" or cleansed[0] == "e2f":

        for z in cleansed[1:]:

            E2F = (translator.translate(z, lang_from='en', lang_to='fr'))
            if "ArgumentOutOfRangeException: 'from' must be a valid language" in E2F.encode('utf-8'):
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            elif E2F.encode('utf-8') == z or E2F.encode('utf-8') == z.capitalize() or E2F.encode('utf-8').capitalize() == z:
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            else:
                trns = (translator.translate(z, lang_from='en', lang_to='fr')).encode('utf-8')
                trns = trns.split()
                trns_lister = []
                for x in trns:
                    x = x.lower()
                    trns_lister.append(x)

                trns_lister = ' '.join(trns_lister)

                if len(trns_lister)>2:
                    final = z + ' ' + '-' + ' ' + str(trns_lister)
                    myarr.append(final)
                else:
                    final = z + ' ' + '-' + ' ' + 'translation was not found'
                    myarr.append(final)


        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:10]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()

        if len(myarr)>10:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[10:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()


        pdf.save()

    if cleansed[0] == "E2G" or cleansed[0] == "e2g":
        for z in cleansed[1:]:
            E2G = (translator.translate(z, lang_from='en', lang_to='de'))
            if "ArgumentOutOfRangeException: 'from' must be a valid language" in E2G.encode('utf-8'):
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            elif E2G.encode('utf-8') == z or E2G.encode('utf-8') == z.capitalize() or E2G.encode('utf-8').capitalize() == z:
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            else:
                trns = (translator.translate(z, lang_from='en', lang_to='de')).encode('utf-8')
                trns = trns.split()
                trns_lister = []
                for x in trns:
                    x = x.lower()
                    trns_lister.append(x)

                trns_lister = ' '.join(trns_lister)

                if len(trns_lister)>2:
                    final = z + ' ' + '-' + ' ' + str(trns_lister)
                    myarr.append(final)
                else:
                    final = z + ' ' + '-' + ' ' + 'translation was not found'
                    myarr.append(final)


        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:10]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()

        if len(myarr)>10:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[10:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()


        pdf.save()

    if cleansed[0] == "E2S" or cleansed[0] == "e2s":
        for z in cleansed[1:]:
            E2S = (translator.translate(z, lang_from='en', lang_to='es'))
            if "ArgumentOutOfRangeException: 'from' must be a valid language" in E2S.encode('utf-8'):
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            elif E2S.encode('utf-8') == z or E2S.encode('utf-8') == z.capitalize() or E2S.encode('utf-8').capitalize() == z:
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            else:
                trns = (translator.translate(z, lang_from='en', lang_to='es')).encode('utf-8')
                trns = trns.split()
                trns_lister = []
                for x in trns:
                    x = x.lower()
                    trns_lister.append(x)

                trns_lister = ' '.join(trns_lister)

                if len(trns_lister)>2:
                    final = z + ' ' + '-' + ' ' + str(trns_lister)
                    myarr.append(final)
                else:
                    final = z + ' ' + '-' + ' ' + 'translation was not found'
                    myarr.append(final)


        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:10]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()

        if len(myarr)>10:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[10:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()


        pdf.save()

    if cleansed[0] == "S2E" or cleansed[0] == "s2e":
        for z in cleansed[1:]:
            S2E = (translator.translate(z, lang_from='es', lang_to='en'))
            if "ArgumentOutOfRangeException: 'from' must be a valid language" in S2E.encode('utf-8'):
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            elif S2E.encode('utf-8') == z or S2E.encode('utf-8') == z.capitalize() or S2E.encode('utf-8').capitalize() == z:
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            else:
                trns = (translator.translate(z, lang_from='es', lang_to='en')).encode('utf-8')
                trns = trns.split()
                trns_lister = []
                for x in trns:
                    x = x.lower()
                    trns_lister.append(x)

                trns_lister = ' '.join(trns_lister)

                if len(trns_lister)>2:
                    final = z + ' ' + '-' + ' ' + str(trns_lister)
                    myarr.append(final)
                else:
                    final = z + ' ' + '-' + ' ' + 'translation was not found'
                    myarr.append(final)


        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:10]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()

        if len(myarr)>10:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[10:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()


        pdf.save()

    if cleansed[0] == "G2E" or cleansed[0] == "g2e":
        for z in cleansed[1:]:

            G2E = (translator.translate(z, lang_from='de', lang_to='en'))
            if "ArgumentOutOfRangeException: 'from' must be a valid language" in G2E.encode('utf-8'):
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            elif G2E.encode('utf-8') == z or G2E.encode('utf-8') == z.capitalize() or G2E.encode('utf-8').capitalize() == z:
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            else:
                trns = (translator.translate(z, lang_from='de', lang_to='en')).encode('utf-8')
                trns = trns.split()
                trns_lister = []
                for x in trns:
                    x = x.lower()
                    trns_lister.append(x)

                trns_lister = ' '.join(trns_lister)
                if len(trns_lister)>2:
                    final = z + ' ' + '-' + ' ' + str(trns_lister)
                    myarr.append(final)
                else:
                    final = z + ' ' + '-' + ' ' + 'translation was not found'
                    myarr.append(final)


        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:10]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()

        if len(myarr)>10:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[10:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()


        pdf.save()

    if cleansed[0] == "F2E" or cleansed[0] == "f2e":
        for z in cleansed[1:]:

            F2E = (translator.translate(z, lang_from='fr', lang_to='en'))

            if "ArgumentOutOfRangeException: 'from' must be a valid language" in F2E.encode('utf-8'):
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)
            elif F2E.encode('utf-8') == z or F2E.encode('utf-8') == z.capitalize() or F2E.encode('utf-8').capitalize() == z:
                final = z + ' ' + '-' + ' ' + 'translation was not found'
                myarr.append(final)

            else:
                trns = (translator.translate(z, lang_from='fr', lang_to='en')).encode('utf-8')
                trns = trns.split()
                trns_lister = []
                for x in trns:
                    x = x.lower()
                    trns_lister.append(x)

                trns_lister = ' '.join(trns_lister)
                if len(trns_lister)>2:
                    final = z + ' ' + '-' + ' ' + str(trns_lister)
                    myarr.append(final)
                else:
                    final = z + ' ' + '-' + ' ' + 'translation was not found'
                    myarr.append(final)


        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:10]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()

        if len(myarr)>10:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[10:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()

        pdf.save()
    if cleansed[0] == "dictionary" or cleansed[0] == "Dictionary":
        for z in cleansed[1:]:
            definition = dictionary.meaning(z)
            if definition != None:
                definer = []
                if 'Noun' in definition:
                    initial = str(definition['Noun'])

                    noun = "noun" + " : " + initial
                    noun = str(noun)
                    definer.append(noun)

                if 'Adjective' in definition:
                    initial = str(definition['Adjective'])

                    adjective = "adjective" + " : " + initial
                    adjective = str(adjective)
                    definer.append(adjective)

                if 'Adverb' in definition:
                    initial = str(definition['Adverb'])

                    adverb = "adverb" + " : " + initial
                    adverb = str(adverb)
                    definer.append(adverb)

                if 'Verb' in definition:
                    initial = str(definition['Verb'])

                    verb = "verb" + " : " + initial
                    verb = str(verb)
                    definer.append(verb)

                if 'Preposition' in definition:
                    initial = str(definition['Preposition'])

                    preposition = "preposition" + " : " + initial
                    preposition = str(preposition)
                    definer.append(preposition)

                if 'Interjection' in definition:
                    initial = str(definition['Interjection'])

                    interjection = "interjection" + " : " + initial
                    interjection = str(interjection)
                    definer.append(interjection)

                final = str(definer)
                final = final.replace("'", "")
                final = final.replace(",", ";")
                final = final.replace("/", "")
                final = final.replace("(", "")
                final = final.replace(")", "")
                final = final.replace("\\", "")

                final = final.replace('"', "")
                final = final.replace("[noun", "noun")
                final = final.replace("[adjective", "adjective")
                final = final.replace("[adverb", "adverb")
                final = final.replace("[verb", "verb")
                final = final.replace("[preposition", "preposition")
                final = final.replace("[interjection", "interjection")
                final = final.replace("]]", "]")
                final = z + '- ' + final
                myarr.append(final)
                print myarr

            else:
                final = z + '-' + 'definition was not found'
                myarr.append(final)

        # output = StringIO()
        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:8]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()



        if len(myarr) > 8:

            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)


            for x in myarr[8:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 11 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue


            pdf.drawText(rhyme)
            pdf.showPage()

        else:
            print

        pdf.save()
    else:
        for z in cleansed:
            results = wiki.find(z)
            #results = [urllib2.unquote(x).decode('utf8') for x in results[:2]]
            article = 'hi'
            try:
                article = wiki.get_article(results[0])
            except IndexError:
                myarr.append(z+' - '+'definition for this term was not found')
                continue
            results = [urllib2.unquote(x.encode('utf8')).decode('utf8') for x in results[:2]]
            print results
            print article.summary.encode('utf8')
            if "may refer to" in article.summary.encode('utf8'):
                print 'True'
                a = results[0]
                b = results [1]
                a = a.replace ("_", " ")
                b = b.replace ("_", " ")
                myarr.append(z+ ' - '+' definition for this term was not found.'+ ' ' + 'did you mean' + ' ' +  a + ' ' + 'or' + ' '  + b + '?')
                continue

            x=unicodedata.normalize("NFKD", article.summary).replace(u'\u0301', '')
            #x = x.replace("O.S.","")
            #x = x.replace("translit.", "")
            #x = x.replace("translit.", "")
            #x = x.replace("translit.", "")
            #x = x.replace("lit.","")
            #x = x.replace("lit.","")
            #x = x.replace("c.","")

            #spade = u"♠"
            #spade2 = spade.encode("utf8")
            #x = x.replace(spade,"")
            #o_accent = u"ō"
            #o_accent2= o_accent.encode("utf8")
            #x = x.replace(o_accent, "")
            #x = re.sub(u"\u2013", "-", x)
            #x = urllib2.unquote(x).decode('utf8')
            #x = x.encode('ascii', 'ignore')
            #x = re.sub(u"\u2013", "-", x)
            final = x
            #final = re.split(r' *[\.\?!][\'"\)\]]* *', final)

            #print type(final)
            final = final.encode('utf8')
            final = re.sub("\xe2\x80\x93", "-", final)

            final = final.decode('utf8')
            #print type(final)


            wordList = ['is a', 'is the', 'was a', 'was the', 'was an', 'is an', 'is one', 'were a', 'were an', 'known as', 'are a', 'are the', 'are an', 'served as']
            startIndex = 0
            try:
                numb = (re.search('|'.join(wordList), x[startIndex:]).start())
                a = numb
                print a
            except:
                a = 0

            #or var.find('are the') or var.find('were the')
            #a=x.find("is")
            title=article.heading
            final=final[a:]
            #final = final.decode('utf8')
            #print final
            print type(final)
            final = final.encode("ascii", "ignore")






            print type(final), 'the final'
            #final = re.sub('\.(?!(\S[^. ])|\d)', '', final)
            #final = re.sub('(?![a-zA-Z]{2}')
            rx = re.compile('\s(?=\.|\?)(?!.\w\.\w)(?!\.[a-z][A-Z])(?!\.[A-Z][ ][a-z]+[A-Z])(?!\.[A-Z][ ]\.[A-Z])')
            final = [x[::-1] for x in rx.split(final[::-1])][::-1]


            #re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', final)
            print 'q', final
            #final = re.split(r' *[\.\?!][\'"\)\]]* *', final)

            if final[1] == ' ':
                final = final[0]+' '+ final[1] + ' ' + final[2]
            else:
                final = final[0]+' '+ final[1]


            print 'z', final
            final = title.encode('utf8') + '   -   ' + final.encode('utf8')


            myarr.append(final)


        #str1 = "".join(myarr)
        #return str1
        #return final



        pdf = Canvas(output, pagesize = letter)
        pdf.setFont('Courier', 40)
        pdf.setFillColorRGB(0, .6, 1)
        pdf.drawCentredString(300, 715, 'Your Study Guide')
        pdf.setFont('Courier', 25)
        pdf.drawCentredString(300, 650, 'Courtesy of StudyGuideMaker.com')
        rhyme = pdf.beginText(inch * .8, inch * 8)
        pdf.setFont('Courier', 12)
        pdf.setFillColorRGB(0, 0, 0)
        counter = 0


        for x in myarr[0:6]:
            wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
            rhyme.textLines(wraped_text)
            rhyme.textLines('\n')


        pdf.drawText(rhyme)
        pdf.showPage()


        if len(myarr) > 6:
            rhyme = pdf.beginText(inch * .8, inch * 10)
            pdf.setFont('Courier', 12)
            pdf.setFillColorRGB(0, 0, 0)

            for x in myarr[6:]:
                wraped_text = "\n".join(wrap(x, 70)) # 80 is line width
                rhyme.textLines(wraped_text)
                rhyme.textLines('\n')
                counter = counter + 1
                if counter % 7 == 0 and counter!=0:
                    pdf.drawText(rhyme)
                    pdf.showPage()
                    rhyme = pdf.beginText(inch * .8, inch * 10)
                    pdf.setFont('Courier', 12)
                    pdf.setFillColorRGB(0, 0, 0)
                    counter = 0
                    continue

            pdf.drawText(rhyme)
            pdf.showPage()

        else:
            pass

        pdf.save()

    pdf_out = output.getvalue()
    temp_dir = "mysite/static/temp/"
    public_temp_path = "/tmp/"
    millis = str(int(round(time.time() * 1000)))
    temp_file_name = millis + '.pdf'
    temp_file_path = temp_dir + temp_file_name
    with open (temp_file_path, 'w') as fd:
        output.seek(0)
        shutil.copyfileobj(output, fd)
    output.close()
    response = make_response(public_temp_path + temp_file_name)
    response.headers['Access-Control-Allow-Origin'] = "*"
    response.headers['Content-Disposition'] = "inline; filename=StudyGuideMaker.pdf"
    response.mimetype = 'application/pdf'
    return response


    #return redirect('/mylink')


@app.route('/helpfultips')
def helpfultips():
    return render_template('helpfultips.html')

@app.route('/SampleStudyGuide')
def SampleStudyGuide():
    return send_file('/home/studyguidemaker/mysite/StudyGuideMaker.pdf', attachment_filename='StudyGuideMaker.pdf', as_attachment=False)

@app.route('/sitemap.xml')
def sitemap():
    return send_file('/home/studyguidemaker/mysite/sitemap.xml', attachment_filename='sitemap.xml', as_attachment=False)

@app.route('/robots.txt')
def robots():
    return send_file('/home/studyguidemaker/mysite/robots.txt', attachment_filename='robots.txt', as_attachment=False)



if __name__ == "__main__":
    app.run()


"""

output = StringIO()

p = canvas.Canvas(output)
p.drawCentredString(300,800, 'Your Study Guide')
for x in myarr:
    counter = 0
    initial = 760
    final = initial - counter
    p.drawString(50, final, x)
    counter = counter + 20
    if final == 0 and counter!=0:
        p.showPage()
        counter = 0
        initial = 760
        final = initial - counter
        p.drawString(50, final, x)


p.showPage()
p.save()

pdf_out = output.getvalue()
output.close()

response = make_response(pdf_out)
response.headers['Content-Disposition'] = "attachment; filename='your_study_guide.pdf"
response.mimetype = 'application/pdf'
return response
"""

"""
document = Document()
document.add_heading("Your Study Guide", 0)

for x in myarr:
    document.add_paragraph(x, style='List Bullet')
f = StringIO()
document.save(f)
length = f.tell()
f.seek(0)

return send_file(f, as_attachment=True, attachment_filename='yourstudyguide.doc')
"""
